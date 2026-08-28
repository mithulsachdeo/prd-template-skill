"""
Converts a structured markdown file (see the DSL documented in SKILL.md) into
the styled PRD .docx used by the prd-template skill.

Usage:
    python generate_docx.py <input.md> <output.docx>

DSL:
    # Title              -> Heading 1
    ## Subtitle           -> Heading 2
    **LABEL TEXT**        -> styled blue bold label (line is ONLY the bold text)
    ⚠ FLAG: ...           -> styled red bold callout
    - bullet text         -> bullet list item
    *italic text*          -> italic paragraph (line is ONLY the italic text)
    | a | b |              -> table row (header row bold; a |---|---| row is the separator)
    \\pagebreak            -> page break
    anything else          -> plain paragraph (leading/trailing ** stripped and bolded if wraps whole line)
"""
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FLAG_COLOR = RGBColor(0xB0, 0x00, 0x00)
LABEL_COLOR = RGBColor(0x1F, 0x4E, 0x79)

LABEL_RE = re.compile(r"^\*\*(.+)\*\*$")
ITALIC_RE = re.compile(r"^\*(.+)\*$")
BULLET_RE = re.compile(r"^-\s+(.*)$")
H1_RE = re.compile(r"^#\s+(.*)$")
H2_RE = re.compile(r"^##\s+(.*)$")
FLAG_RE = re.compile(r"^⚠\s*FLAG:\s*(.*)$", re.IGNORECASE)
TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
TABLE_SEP_RE = re.compile(r"^\|[\s:|-]+\|$")


def parse_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_label(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = LABEL_COLOR
    r.font.size = Pt(10.5)


def add_flag(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("⚠ FLAG: " + text)
    r.bold = True
    r.font.color.rgb = FLAG_COLOR


def add_table(doc, rows):
    if not rows:
        return
    t = doc.add_table(rows=1, cols=len(rows[0]))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, val in enumerate(rows[0]):
        hdr[i].text = val
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows[1:]:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            if i < len(cells):
                cells[i].text = val
    doc.add_paragraph()


def convert(input_path, output_path):
    with open(input_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    table_buffer = []

    def flush_table():
        if table_buffer:
            add_table(doc, list(table_buffer))
            table_buffer.clear()

    while i < n:
        raw = lines[i]
        line = raw.strip()

        if TABLE_ROW_RE.match(line):
            if TABLE_SEP_RE.match(line):
                i += 1
                continue
            table_buffer.append(parse_table_row(line))
            i += 1
            continue
        else:
            flush_table()

        if not line:
            i += 1
            continue

        if line == "\\pagebreak":
            doc.add_page_break()
            i += 1
            continue

        m = H1_RE.match(line)
        if m:
            doc.add_heading(m.group(1), level=1)
            i += 1
            continue

        m = H2_RE.match(line)
        if m:
            doc.add_heading(m.group(1), level=2)
            i += 1
            continue

        m = FLAG_RE.match(line)
        if m:
            add_flag(doc, m.group(1))
            i += 1
            continue

        m = LABEL_RE.match(line)
        if m:
            add_label(doc, m.group(1))
            i += 1
            continue

        m = ITALIC_RE.match(line)
        if m:
            p = doc.add_paragraph()
            r = p.add_run(m.group(1))
            r.italic = True
            i += 1
            continue

        m = BULLET_RE.match(line)
        if m:
            doc.add_paragraph(m.group(1), style="List Bullet")
            i += 1
            continue

        # Centered title convention: a line wrapped in [[ ]] is centered plain text
        if line.startswith("[[") and line.endswith("]]"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line[2:-2])
            i += 1
            continue

        doc.add_paragraph(line)
        i += 1

    flush_table()
    doc.save(output_path)
    print(f"saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
